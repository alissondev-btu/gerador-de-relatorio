import pandas as pd  # Biblioteca para manipulação de dados em planilhas Excel
from docx import Document  # Biblioteca para manipulação de documentos Word (.docx)
import tkinter as tk  # Biblioteca para criar interface gráfica
from tkinter import filedialog, messagebox  # Componentes para diálogo de arquivos e mensagens na interface
import os  # Biblioteca para manipulação de arquivos e diretórios

def substituir_placeholders(documento, row, colunas):
    # Substitui os placeholders no documento Word pelos dados da linha da planilha
    
    # Substituição nos parágrafos do documento
    for paragrafo in documento.paragraphs:
        for col in colunas:
            placeholder = f'{{{{{col}}}}}'  # Cria o texto do placeholder, ex: {{nome_cliente}}
            if placeholder in paragrafo.text:
                # Substitui o placeholder pelo valor da coluna correspondente naquela linha
                paragrafo.text = paragrafo.text.replace(placeholder, str(row[col]))
    
    # Substituição dentro das tabelas do documento
    for table in documento.tables:
        for row_tables in table.rows:
            for cell in row_tables.cells:
                for col in colunas:
                    placeholder = f'{{{{{col}}}}}'
                    if placeholder in cell.text:
                        # Substitui o placeholder no texto da célula da tabela
                        cell.text = cell.text.replace(placeholder, str(row[col]))

def gerar_documentos():
    try:
        # Abre janela para o usuário selecionar o arquivo Excel (.xlsx)
        excel_path = filedialog.askopenfilename(title="Selecione o arquivo Excel", filetypes=[("Arquivos Excel", "*.xlsx")])
        if not excel_path:
            return  # Se o usuário cancelar, sai da função

        # Abre janela para selecionar o arquivo modelo Word (.docx)
        modelo_path = filedialog.askopenfilename(title="Selecione o modelo Word", filetypes=[("Documentos Word", "*.docx")])
        if not modelo_path:
            return

        # Lê os dados do Excel para um DataFrame pandas
        df = pd.read_excel(excel_path)

        # Abre janela para selecionar a pasta onde os relatórios serão salvos
        pasta_saida = filedialog.askdirectory(title="Selecione a pasta de saída")
        if not pasta_saida:
            return

        # Para cada linha da planilha, gera um documento Word personalizado
        for idx, row in df.iterrows():
            documento = Document(modelo_path)  # Abre o modelo Word
            substituir_placeholders(documento, row, df.columns)  # Substitui os placeholders pelos dados da linha

            # Cria o nome do arquivo baseado na coluna 'nome_cliente'
            nome_cliente = str(row['nome_cliente']).strip().replace(' ', '_')
            nome_arquivo = os.path.join(pasta_saida, f"relatorio_{nome_cliente}.docx")
            documento.save(nome_arquivo)  # Salva o documento na pasta escolhida

        # Mostra mensagem de sucesso ao final do processo
        messagebox.showinfo("Sucesso", "Relatórios gerados com sucesso!")

    except Exception as e:
        # Se ocorrer qualquer erro, mostra uma mensagem de erro com a descrição
        messagebox.showerror("Erro", f"Ocorreu um erro: {str(e)}")

# Criação da interface gráfica principal com Tkinter
janela = tk.Tk()
janela.title("Gerador de Relatórios Word")  # Título da janela
janela.geometry("300x200")  # Tamanho da janela

# Texto explicativo na janela
rotulo = tk.Label(janela, text="Gerador de Relatórios Word\ncom base em planilha Excel", font=("Arial", 12), justify="center")
rotulo.pack(pady=20)  # Posiciona o texto com espaçamento vertical

# Botão que inicia o processo de geração dos relatórios
botao_gerar = tk.Button(janela, text="Selecionar arquivos e Gerar", command=gerar_documentos, bg="#4CAF50", fg="white", font=("Arial", 11))
botao_gerar.pack(pady=10)  # Posiciona o botão com espaçamento vertical

# Mantém a janela aberta, aguardando interações do usuário
janela.mainloop()
